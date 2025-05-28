"""
FastAPI application for Resume Analyzer.
"""
import os
import json
from typing import Dict, Any, Optional
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
import uvicorn
import mlflow
from io import BytesIO
import time
import os
from dotenv import load_dotenv
from pathlib import Path
import logging
from fastapi import status

load_dotenv()

from app.config import APP_NAME, APP_VERSION, DEBUG, MLFLOW_TRACKING_URI, EXPERIMENT_NAME
from app.core.security import (
    verify_password, get_password_hash, create_access_token,
    verify_token, Token, TokenData
)
from app.models.schema import (
    ResumeAnalysisRequest, ResumeAnalysisResponse,
    InterviewQuestionRequest, InterviewQuestionsResponse, MockInterviewFeedbackRequest, MockInterviewFeedbackResponse,
    SalaryIntelligenceRequest, SalaryIntelligenceResponse
)
from app.agents.react_agent import ResumeReactAgent
from app.services.parser import extract_skills, extract_experience, extract_education
from app.services.analyzer import analyze_strengths_weaknesses, calculate_job_match, suggest_improvements
from app.services.vector_store import initialize_vector_store, get_similar_skills
from app.services.logging import initialize_promptlayer
from app.services.structured_analyzer import StructuredAnalyzer
from app.services.resume_builder import ResumeBuilder
from app.routers import career_paths # Import only career_paths for now

# Setup logging
logger = logging.getLogger(__name__)

# Initialize the FastAPI app
app = FastAPI(
    title=APP_NAME,
    version=APP_VERSION,
    description="An AI-powered resume analyzer using ReAct agents",
    docs_url="/docs",
    redoc_url="/redoc",
    debug=DEBUG
)

# Set up templates and static files
templates = Jinja2Templates(directory="app/templates")
app.mount("/static", StaticFiles(directory="app/static"), name="static")

# Set up CORS with proper configuration
origins = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["*"],
    max_age=3600,
)

# OAuth2 scheme for token authentication
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

# Add security dependencies
async def get_current_user(token: str = Depends(oauth2_scheme)):
    return verify_token(token)

# Add login endpoint
@app.post("/token", response_model=Token)
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends()):
    # In a real application, you would validate against a database
    # This is just a placeholder
    if form_data.username != "test" or not verify_password(form_data.password, get_password_hash("test")):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token = create_access_token(data={"sub": form_data.username})
    return {"access_token": access_token, "token_type": "bearer"}

# Initialize MLflow
mlflow.set_tracking_uri(MLFLOW_TRACKING_URI)
mlflow.set_experiment(EXPERIMENT_NAME)

# Initialize analyzers
agent = None
structured_analyzer = StructuredAnalyzer()
resume_builder = ResumeBuilder()

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    """Render the main page."""
    import time
    import random
    # Add cache busting parameter and version
    cache_buster = int(time.time())
    version = f"v3.{random.randint(1000, 9999)}"  # Random version to force cache refresh
    
    # Add no-cache headers
    response = templates.TemplateResponse("index.html", {
        "request": request, 
        "cache_buster": cache_buster,
        "version": version
    })
    
    # Force no caching
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    
    return response

@app.get("/resume-builder", response_class=HTMLResponse)
async def resume_builder_page(request: Request):
    """Render the resume builder page."""
    return templates.TemplateResponse("resume_builder.html", {"request": request})

@app.get("/interview-assistant", response_class=HTMLResponse)
async def interview_assistant_page(request: Request):
    """Render the interview assistant page."""
    return templates.TemplateResponse("interview_assistant.html", {"request": request})

@app.get("/salary-intelligence")
async def salary_intelligence(request: Request):
    """Serve the Salary Intelligence page."""
    return templates.TemplateResponse("salary_intelligence.html", {"request": request})

@app.get("/resume-analysis")
async def resume_analysis(request: Request):
    """Serve the Resume Analysis page."""
    return templates.TemplateResponse("resume_analysis.html", {"request": request})

@app.get("/skills-career-dashboard", response_class=HTMLResponse)
async def skills_career_dashboard_page(request: Request):
    """Render the skills and career dashboard page."""
    return templates.TemplateResponse("skills_career_dashboard.html", {"request": request})

# Initialize services on startup
@app.on_event("startup")
async def startup_event():
    # Initialize PromptLayer
    initialize_promptlayer()
    
    # Initialize vector store
    # initialize_vector_store()  # Commented out due to ChromaDB schema issue
    
    # Initialize the agent (this will be done on-demand, but we can test it here)
    try:
        global agent
        agent = ResumeReactAgent()
    except Exception as e:
        print(f"Failed to initialize agent: {str(e)}")
        # Continue anyway, we'll initialize on-demand

def get_agent():
    """Get or create the ReAct agent."""
    global agent
    if agent is None:
        agent = ResumeReactAgent()
    return agent

# Removed duplicate root endpoint - using the HTML template version above

@app.get("/health")
async def health_check():
    """Health check endpoint."""
    return {
        "status": "healthy",
        "message": "Server is running properly",
        "timestamp": "2025-01-26"
    }

@app.post("/analyze/text", response_model=Dict[str, Any])
async def analyze_resume_text(
    request: ResumeAnalysisRequest,
    current_user: TokenData = Depends(get_current_user)
):
    """
    Analyze a resume from text input.
    
    Args:
        request: ResumeAnalysisRequest with resume_text and optional job_description
        
    Returns:
        Analysis results
    """
    # Start MLflow run
    with mlflow.start_run() as run:
        mlflow.log_param("resume_length", len(request.resume_text))
        mlflow.log_param("has_job_description", request.job_description is not None)
        
        start_time = time.time()
        
        try:
            # Get the agent
            agent = get_agent()
            
            # Execute the agent
            result = agent.analyze_resume(
                resume_text=request.resume_text,
                job_description=request.job_description
            )
            
            # Log metrics
            processing_time = time.time() - start_time
            mlflow.log_metric("processing_time_seconds", processing_time)
            mlflow.log_metric("success", 1)
            
            return result
            
        except Exception as e:
            # Log failure
            mlflow.log_metric("success", 0)
            mlflow.log_param("error", str(e))
            
            raise HTTPException(
                status_code=500,
                detail=f"Failed to analyze resume: {str(e)}"
            )

@app.post("/analyze/file", response_model=Dict[str, Any])
async def analyze_resume_file(
    file: UploadFile = File(...),
    job_description: Optional[str] = Form(None),
    current_user: TokenData = Depends(get_current_user)
):
    """
    Analyze a resume from a file upload.
    
    Args:
        file: The resume file (PDF or DOCX)
        job_description: Optional job description for matching
        
    Returns:
        Analysis results
    """
    try:
        # Read the file content
        content = await file.read()
        
        # Parse the resume based on file type
        resume_text = ""
        
        if file.filename.endswith(".pdf"):
            from pypdf import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(content))
            for page in reader.pages:
                resume_text += page.extract_text()
                
        elif file.filename.endswith(".docx"):
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(content))
            resume_text = "\n".join([p.text for p in doc.paragraphs])
                
        else:
            # Assume it's plain text
            resume_text = content.decode("utf-8")
        
        # Get the agent
        agent = get_agent()
        
        # Execute the agent
        result = agent.analyze_resume(
            resume_text=resume_text,
            job_description=job_description
        )
        
        return result
            
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to analyze resume: {str(e)}"
        )

@app.post("/simple_analyze", response_model=Dict[str, Any])
async def simple_analyze_resume(
    file: UploadFile = File(...),
    job_description: Optional[str] = Form(None)
):
    """
    A simplified version of resume analysis that uses fewer API calls.
    """
    try:
        # Read the file content
        content = await file.read()
        
        # Parse the resume based on file type
        resume_text = ""
        
        if file.filename.endswith(".pdf"):
            from pypdf import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(content))
            for page in reader.pages:
                resume_text += page.extract_text()
                
        elif file.filename.endswith(".docx"):
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(content))
            resume_text = "\n".join([p.text for p in doc.paragraphs])
                
        else:
            # Assume it's plain text
            resume_text = content.decode("utf-8")
        
        # Get the agent
        agent = get_agent()
        
        # Use the direct analyze method to avoid agent overhead
        result = agent.direct_analyze(
            resume_text=resume_text,
            job_description=job_description
        )
        
        return result
            
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to analyze resume: {str(e)}"
        )

@app.post("/structured_analyze", response_model=Dict[str, Any])
async def structured_analyze(
    file: UploadFile = File(...),
    job_description: Optional[str] = Form(None)
):
    """
    Analyze a resume with structured output that doesn't use API calls.
    """
    try:
        # Read the file content
        content = await file.read()
        
        # Parse the resume based on file type
        resume_text = ""
        
        if file.filename.endswith(".pdf"):
            from pypdf import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(content))
            for page in reader.pages:
                resume_text += page.extract_text()
                
        elif file.filename.endswith(".docx"):
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(content))
            resume_text = "\n".join([p.text for p in doc.paragraphs])
                
        else:
            # Assume it's plain text
            resume_text = content.decode("utf-8")
        
        # Use the structured analyzer
        result = structured_analyzer.analyze_resume(
            resume_text=resume_text,
            job_description=job_description
        )
        
        return result
            
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to analyze resume: {str(e)}"
        )

@app.post("/test_analyze")
async def test_analyze():
    """Test endpoint to verify response format."""
    return {
        "analysis": "This is a test analysis response. The candidate shows strong technical skills in Python and machine learning.",
        "skill_match_details": {
            "match_percentage": 85.5,
            "matched_skills": ["Python", "Machine Learning", "TensorFlow"],
            "resume_skills": ["Python", "Machine Learning", "TensorFlow", "AWS", "Docker"],
            "job_description_skills": ["Python", "Machine Learning", "TensorFlow", "React"]
        },
        "success": True
    }

@app.get("/test_analyze")
async def test_analyze_get():
    """Test endpoint for connection testing (GET version)."""
    return {
        "status": "success",
        "message": "Connection test successful",
        "timestamp": time.time(),
        "method": "GET"
    }

@app.post("/test-file-upload")
async def test_file_upload(
    file: UploadFile = File(...),
    job_description: str = Form(...),
    focus_areas: Optional[str] = Form(None)
):
    """Test endpoint to verify file upload is working."""
    try:
        content = await file.read()
        return {
            "success": True,
            "filename": file.filename,
            "file_size": len(content),
            "content_type": file.content_type,
            "job_description_length": len(job_description),
            "focus_areas": focus_areas
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

@app.post("/analyze")
async def analyze(file: UploadFile = File(...), job_description: Optional[str] = Form(None)):
    try:
        # Validate file type
        if not file.filename.endswith(('.pdf', '.docx', '.txt')):
            raise HTTPException(status_code=400, detail="File type not supported. Please upload a PDF, DOCX, or TXT file.")
        
        # Read the file content
        content = await file.read()
        
        # Parse the resume based on file type
        resume_text = ""
        
        if file.filename.endswith(".pdf"):
            from pypdf import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(content))
            for page in reader.pages:
                resume_text += page.extract_text()
                
        elif file.filename.endswith(".docx"):
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(content))
            resume_text = "\n".join([p.text for p in doc.paragraphs])
                
        else:
            # Assume it's plain text
            resume_text = content.decode("utf-8")
        
        # Get the agent
        agent = get_agent()
        
        # Execute the agent
        result = agent.analyze_resume(
            resume_text=resume_text,
            job_description=job_description
        )
        
        return result
            
    except Exception as e:
        logger.error(f"Error analyzing resume: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# Resume Builder Endpoints
@app.post("/build-resume")
async def build_resume(
    name: str = Form(...),
    email: str = Form(...),
    phone: str = Form(...),
    skills: str = Form(...),
    experience: str = Form(...),
    education: str = Form(...),
    job_description: Optional[str] = Form(None),
    target_role: Optional[str] = Form(None)
):
    """
    Build an optimized resume from user information.
    """
    try:
        # Structure user information
        user_info = {
            "personal_info": {
                "name": name,
                "email": email,
                "phone": phone
            },
            "skills": [skill.strip() for skill in skills.split(',') if skill.strip()],
            "experience": experience,
            "education": education
        }
        
        # Build the resume
        result = resume_builder.build_resume(
            user_info=user_info,
            job_description=job_description,
            target_role=target_role
        )
        
        return result
        
    except Exception as e:
        logger.error(f"Error building resume: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/optimize-resume")
async def optimize_resume(
    file: Optional[UploadFile] = File(None),
    resume_text: Optional[str] = Form(None),
    job_description: str = Form(...),
    focus_areas: Optional[str] = Form(None)
):
    """
    Optimize an existing resume for a specific job.
    Can accept either a file upload or direct text input.
    """
    resume_text_content = ""
    used_input_type = ""

    try:
        logger.info(f"Optimize request received. Job description length: {len(job_description)}, Focus areas: {focus_areas}")

        if file and file.filename and file.size > 0:
            logger.info(f"Processing uploaded file: {file.filename}, size: {file.size}, content_type: {file.content_type}")
            used_input_type = f"file: {file.filename}"
            # Read the file content
            content = await file.read()
            logger.info(f"File content read successfully, size: {len(content)} bytes")
            
            # Parse the resume
            if file.filename.endswith(".pdf"):
                from pypdf import PdfReader
                from io import BytesIO
                reader = PdfReader(BytesIO(content))
                for page in reader.pages:
                    resume_text_content += page.extract_text()
            elif file.filename.endswith(".docx"):
                import docx
                from io import BytesIO
                doc = docx.Document(BytesIO(content))
                resume_text_content = "\n".join([p.text for p in doc.paragraphs])
            elif file.filename.endswith(".txt"):
                resume_text_content = content.decode("utf-8")
            else:
                logger.warning(f"Unsupported file type for optimization: {file.filename}")
                raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.filename}. Please use PDF, DOCX, or TXT.")
        
        elif resume_text and resume_text.strip():
            logger.info(f"Processing resume_text from form, length: {len(resume_text)}")
            used_input_type = "text_input"
            resume_text_content = resume_text.strip()
        
        else:
            logger.warning("No resume content provided (neither file nor text).")
            raise HTTPException(status_code=400, detail="No resume content provided. Please upload a file or paste resume text.")

        if not resume_text_content.strip():
            logger.warning(f"Resume content is empty after attempting to process from {used_input_type}.")
            raise HTTPException(status_code=400, detail="Resume content is empty. Please provide valid content.")

        logger.info(f"Resume text extracted for optimization (from {used_input_type}), length: {len(resume_text_content)}")
        
        # Parse focus areas
        focus_list = None
        if focus_areas:
            focus_list = [area.strip() for area in focus_areas.split(',') if area.strip()]
        
        # Optimize the resume
        result = resume_builder.optimize_existing_resume(
            current_resume=resume_text_content,
            job_description=job_description,
            focus_areas=focus_list
        )
        
        logger.info(f"Optimization completed successfully: {result.get('success', False)}")
        return result
        
    except Exception as e:
        logger.error(f"Error optimizing resume: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-resume-versions")
async def generate_resume_versions(
    name: str = Form(...),
    email: str = Form(...),
    phone: str = Form(...),
    skills: str = Form(...),
    experience: str = Form(...),
    education: str = Form(...),
    job_description: str = Form(...),
    num_versions: int = Form(3)
):
    """
    Generate multiple versions of a resume for A/B testing.
    """
    try:
        # Structure user information
        user_info = {
            "personal_info": {
                "name": name,
                "email": email,
                "phone": phone
            },
            "skills": [skill.strip() for skill in skills.split(',') if skill.strip()],
            "experience": experience,
            "education": education
        }
        
        # Generate multiple versions
        result = resume_builder.generate_multiple_versions(
            user_info=user_info,
            job_description=job_description,
            num_versions=min(num_versions, 3)  # Limit to 3 versions
        )
        
        return {"versions": result, "success": True}
        
    except Exception as e:
        logger.error(f"Error generating resume versions: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/interview/generate-questions", response_model=InterviewQuestionsResponse)
async def generate_interview_questions(request: InterviewQuestionRequest):
    """
    Generate personalized interview questions based on resume and job description.
    """
    try:
        # Get the agent
        agent = get_agent()
        
        # Generate questions using the AI agent
        result = agent.generate_interview_questions(
            resume_text=request.resume_text,
            job_description=request.job_description,
            question_types=request.question_types,
            num_questions=request.num_questions
        )
        
        return result
        
    except Exception as e:
        # Return error response with fallback questions
        error_questions = [
            InterviewQuestion(
                question=f"Error generating questions: {str(e)}",
                type="error"
            )
        ]
        return InterviewQuestionsResponse(questions=error_questions)

@app.post("/interview/mock-interview-feedback", response_model=MockInterviewFeedbackResponse)
async def mock_interview_feedback(request: MockInterviewFeedbackRequest):
    """Provide AI feedback on mock interview answers."""
    try:
        agent = ResumeReactAgent()
        feedback = agent.get_mock_interview_feedback(
            question=request.question,
            answer=request.answer,
            job_description=request.job_description
        )
        return feedback
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error providing interview feedback: {str(e)}")

@app.post("/salary/analyze", response_model=SalaryIntelligenceResponse)
async def analyze_salary_intelligence(request: SalaryIntelligenceRequest):
    """Analyze salary intelligence and provide market insights."""
    try:
        agent = ResumeReactAgent()
        analysis = agent.analyze_salary_intelligence(
            resume_text=request.resume_text,
            job_title=request.job_title,
            location=request.location,
            years_of_experience=request.years_of_experience,
            company_size=request.company_size,
            industry=request.industry
        )
        return analysis
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error analyzing salary intelligence: {str(e)}")

@app.post("/extract-resume-text")
async def extract_resume_text(file: UploadFile = File(...)):
    """
    Extract text from uploaded resume file for interview assistant.
    """
    try:
        # Read the file content
        content = await file.read()
        
        # Parse the resume based on file type
        resume_text = ""
        
        if file.filename.endswith(".pdf"):
            from pypdf import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(content))
            for page in reader.pages:
                resume_text += page.extract_text()
                
        elif file.filename.endswith(".docx"):
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(content))
            resume_text = "\n".join([p.text for p in doc.paragraphs])
                
        else:
            # Assume it's plain text
            resume_text = content.decode("utf-8")
        
        return {
            "success": True,
            "resume_text": resume_text,
            "filename": file.filename
        }
            
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "resume_text": ""
        }

@app.post("/analyze-resume", response_model=ResumeAnalysisResponse)
async def analyze_resume(file: UploadFile = File(...), job_description: Optional[str] = Form(None)):
    try:
        # Validate file type
        if not file.filename.endswith(('.pdf', '.docx', '.txt')):
            raise HTTPException(status_code=400, detail="File type not supported. Please upload a PDF, DOCX, or TXT file.")
        
        # Read the file content
        content = await file.read()
        
        # Parse the resume based on file type
        resume_text = ""
        
        if file.filename.endswith(".pdf"):
            from pypdf import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(content))
            for page in reader.pages:
                resume_text += page.extract_text()
                
        elif file.filename.endswith(".docx"):
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(content))
            resume_text = "\n".join([p.text for p in doc.paragraphs])
                
        else:
            # Assume it's plain text
            resume_text = content.decode("utf-8")
        
        # Get the agent
        agent = get_agent()
        
        # Execute the agent
        result = agent.analyze_resume(
            resume_text=resume_text,
            job_description=job_description
        )
        
        return result
            
    except Exception as e:
        logger.error(f"Error analyzing resume: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

app.include_router(career_paths.router, prefix="/api/v1/career-paths", tags=["Skills & Career Paths"])

# Placeholder for WebSocket endpoint for real-time notifications
# @app.websocket("/ws/notifications/{user_id}")

# Add security headers middleware
@app.middleware("http")
async def add_security_headers(request: Request, call_next):
    response = await call_next(request)
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-XSS-Protection"] = "1; mode=block"
    response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    response.headers["Content-Security-Policy"] = "default-src 'self'; script-src 'self' 'unsafe-inline' 'unsafe-eval'; style-src 'self' 'unsafe-inline';"
    return response

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", 8000)))